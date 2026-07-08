"""Generate the Dashboard Setup Guide as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(SCRIPT_DIR, "Dashboard Setup Guide.docx")

doc = Document()

# ── Styles ──────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level, size, color in [
    ("Heading 1", 22, RGBColor(0x1A, 0x23, 0x32)),
    ("Heading 2", 16, RGBColor(0x21, 0x96, 0xF3)),
    ("Heading 3", 13, RGBColor(0x33, 0x33, 0x33)),
]:
    s = doc.styles[level]
    s.font.name = "Calibri"
    s.font.size = Pt(size)
    s.font.color.rgb = color
    s.font.bold = True
    s.paragraph_format.space_before = Pt(18)
    s.paragraph_format.space_after = Pt(8)


def add_para(text, bold=False, italic=False, size=None, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    return p


def add_numbered(text, level=0):
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    return p


def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return p


def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    run._element.get_or_add_rPr().append(shd)
    return p


# ══════════════════════════════════════════════════════════
# COVER / TITLE
# ══════════════════════════════════════════════════════════
doc.add_paragraph()  # spacing
doc.add_paragraph()
add_para("UrbanCo Stock Dashboard", bold=True, size=28,
         color=RGBColor(0x1A, 0x23, 0x32), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Setup & User Guide", bold=True, size=18,
         color=RGBColor(0x21, 0x96, 0xF3), align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para("A step-by-step guide for setting up, viewing, and maintaining\n"
         "the daily auto-updating stock dashboard.", size=12,
         color=RGBColor(0x66, 0x66, 0x66), align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para("April 2026", size=11,
         color=RGBColor(0x99, 0x99, 0x99), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. What is the Dashboard?",
    "2. How to View the Dashboard",
    "3. What You'll See on the Dashboard",
    "4. How It Updates Automatically",
    "5. How to Manually Refresh Data",
    "6. How to Set Up from Scratch (For New Users)",
    "7. How to Add or Remove Stocks",
    "8. Troubleshooting",
    "9. Quick Reference Card",
]
for item in toc_items:
    add_para(item, size=11, color=RGBColor(0x21, 0x96, 0xF3))

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# SECTION 1
# ══════════════════════════════════════════════════════════
doc.add_heading("1. What is the Dashboard?", level=1)

doc.add_paragraph(
    "The UrbanCo Stock Dashboard is a single web page that shows you "
    "live stock data for Urban Company and peer companies. It updates "
    "automatically every weekday evening after market close."
)

doc.add_heading("What it tracks:", level=3)
add_bullet("Volume (NSE + BSE combined) with daily, weekly, and monthly averages")
add_bullet("Delivery volume and delivery % for each time period")
add_bullet("Daily price movement (close, high-low range, % change)")
add_bullet("Peer comparison table: Urbanco, Eternal, Swiggy, Meesho, Groww, "
           "Lenskart, Physics Wallah, Blackbuck, Pinelabs")
add_bullet("Stock exchange intimations / corporate announcements")
add_bullet("Results date markers on all charts")

# ══════════════════════════════════════════════════════════
# SECTION 2
# ══════════════════════════════════════════════════════════
doc.add_heading("2. How to View the Dashboard", level=1)

doc.add_paragraph("Simply open this link in any web browser (Chrome, Edge, Safari):")
doc.add_paragraph()
add_para("https://ashishbansal-uc.github.io/urbanco-dashboard/dashboard.html",
         bold=True, size=12, color=RGBColor(0x21, 0x96, 0xF3))
doc.add_paragraph()

doc.add_paragraph(
    "No login, no account, no software needed. Anyone with the link can view it."
)

doc.add_heading("Tip: Bookmark it", level=3)
add_numbered("Open the link above in your browser")
add_numbered("Press Ctrl+D (Windows) or Cmd+D (Mac) to bookmark it")
add_numbered("Check it anytime after 8:30 PM on trading days for fresh data")

add_note("Note: If the page looks outdated, press Ctrl+Shift+R to force-refresh "
         "and clear your browser cache.")

# ══════════════════════════════════════════════════════════
# SECTION 3
# ══════════════════════════════════════════════════════════
doc.add_heading("3. What You'll See on the Dashboard", level=1)

doc.add_heading("Header Bar", level=2)
doc.add_paragraph(
    "Shows the latest close price, daily % change, day high/low, total volume, "
    "delivery %, and 52-week high/low."
)

doc.add_heading("Volume & Delivery Chart", level=2)
doc.add_paragraph(
    "A combined NSE + BSE chart with three tabs you can click:"
)

# Table for volume tabs
table = doc.add_table(rows=4, cols=2)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0]
hdr.cells[0].text = "Tab"
hdr.cells[1].text = "What It Shows"
for cell in hdr.cells:
    for p in cell.paragraphs:
        p.runs[0].bold = True

data = [
    ("Daily", "Daily total volume (stacked bars) + delivery volume (green line) + delivery % (dashed line, right axis)"),
    ("Weekly Avg", "5-day rolling average of total volume, delivery volume, and delivery %"),
    ("Monthly Avg", "22-day rolling average of total volume, delivery volume, and delivery %"),
]
for i, (tab, desc) in enumerate(data):
    table.rows[i + 1].cells[0].text = tab
    table.rows[i + 1].cells[1].text = desc

doc.add_paragraph()
add_note("Orange dashed vertical lines on the charts mark actual results/board meeting dates "
         "as announced on NSE.")

doc.add_heading("Price Movement Chart", level=2)
doc.add_paragraph("Three views via tabs:")
add_bullet("Close Price - daily closing price line chart")
add_bullet("High-Low Range - shows the trading range each day with close overlaid")
add_bullet("Daily Change % - green/red bars showing % up or down each day")

doc.add_heading("Peer Performance Table", level=2)
doc.add_paragraph(
    "Compares 9 stocks across multiple time periods: 1 Day, 1 Week, 30 Days, "
    "90 Days, 180 Days, and 365 Days. Green = gain, Red = loss. "
    "UrbanCo's row is highlighted. Also shows current market price (CMP) and "
    "market cap in thousands of crores."
)

doc.add_heading("Stock Exchange Intimations", level=2)
doc.add_paragraph(
    "Lists the most recent 20 corporate announcements filed by Urban Company "
    "with NSE, including board meeting outcomes, press releases, ESOP updates, "
    "and analyst meeting notifications."
)

# ══════════════════════════════════════════════════════════
# SECTION 4
# ══════════════════════════════════════════════════════════
doc.add_heading("4. How It Updates Automatically", level=1)

doc.add_paragraph(
    "The dashboard is hosted on GitHub Pages and uses GitHub Actions "
    "(a free automation tool) to refresh data every weekday."
)

doc.add_heading("Schedule:", level=3)

table2 = doc.add_table(rows=4, cols=2)
table2.style = "Light Grid Accent 1"
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
rows = [
    ("When", "Every Monday to Friday at 8:30 PM IST"),
    ("Why 8:30 PM?", "NSE/BSE publish final volume and delivery data (bhav copy) by ~8:00 PM"),
    ("On holidays?", "The script runs but there's no new market data, so the dashboard keeps showing the last trading day's data"),
]
table2.rows[0].cells[0].text = "Question"
table2.rows[0].cells[1].text = "Answer"
for cell in table2.rows[0].cells:
    for p in cell.paragraphs:
        p.runs[0].bold = True
for i, (q, a) in enumerate(rows):
    table2.rows[i + 1].cells[0].text = q
    table2.rows[i + 1].cells[1].text = a

doc.add_paragraph()
add_note("You don't need to do anything. It's fully automatic. Just open the link "
         "and you'll always see the latest data.")

# ══════════════════════════════════════════════════════════
# SECTION 5
# ══════════════════════════════════════════════════════════
doc.add_heading("5. How to Manually Refresh Data", level=1)

doc.add_paragraph(
    "If you want to refresh the data right now (without waiting for 8:30 PM), "
    "you have two options:"
)

doc.add_heading("Option A: Trigger from GitHub (recommended)", level=2)
add_numbered('Go to: https://github.com/ashishbansal-UC/urbanco-dashboard/actions')
add_numbered('Click on "Update Dashboard Data" in the left sidebar')
add_numbered('Click the "Run workflow" button on the right')
add_numbered('Click the green "Run workflow" button')
add_numbered("Wait 2-3 minutes, then refresh the dashboard URL")

doc.add_heading("Option B: Run locally on your PC", level=2)
add_numbered('Open the folder: C:\\Users\\Ashish Bansal\\Desktop\\Claude work\'\\Output\\UrbanCo Dashboard')
add_numbered("Double-click refresh.bat")
add_numbered("Wait for it to finish (takes about 1 minute)")
add_numbered("Open dashboard.html in your browser")

add_note("Option B only updates your local copy. To update the live URL for everyone, "
         "you'd also need to push to GitHub (Option A is easier).")

# ══════════════════════════════════════════════════════════
# SECTION 6
# ══════════════════════════════════════════════════════════
doc.add_heading("6. How to Set Up from Scratch (For New Users)", level=1)

doc.add_paragraph(
    "If someone else wants to create a similar dashboard for a different stock, "
    "here's the full setup process:"
)

doc.add_heading("Prerequisites", level=2)
add_bullet("A computer with Python installed (python.org - free)")
add_bullet("A GitHub account (github.com - free)")
add_bullet("Internet connection")

doc.add_heading("Step-by-Step Setup", level=2)

add_numbered("Install Python packages: Open a terminal/command prompt and run:")
add_code("pip install yfinance pandas requests")

add_numbered("Get the dashboard files: Download or copy these files into a folder:")
add_bullet("fetch_data.py - fetches stock data from NSE/BSE", level=1)
add_bullet("dashboard.html - the visual dashboard", level=1)
add_bullet("requirements.txt - package list for GitHub Actions", level=1)
add_bullet("refresh.bat - manual refresh shortcut", level=1)
add_bullet(".github/workflows/update-dashboard.yml - auto-update config", level=1)

add_numbered("Edit the stock symbol in fetch_data.py:")
add_code('STOCK_NSE = "YOURSTOCKNAME.NS"')
add_code('STOCK_BSE = "YOURSTOCKNAME.BO"')
add_code('NSE_SYMBOL = "YOURSTOCKNAME"')

add_numbered("Run it once to test:")
add_code("python fetch_data.py")

add_numbered("Open dashboard.html in your browser to verify it works")

add_numbered("Create a GitHub repository:")
add_bullet("Go to github.com/new", level=1)
add_bullet("Name it (e.g., my-stock-dashboard)", level=1)
add_bullet("Keep it Public", level=1)
add_bullet("Don't add README or .gitignore", level=1)

add_numbered("Push files to GitHub:")
add_code("git init && git branch -m main")
add_code("git add -A && git commit -m 'Initial commit'")
add_code("git remote add origin https://github.com/YOURUSERNAME/REPONAME.git")
add_code("git push -u origin main")

add_numbered("Enable GitHub Pages:")
add_bullet("Go to repo Settings > Pages", level=1)
add_bullet('Set Source to "Deploy from a branch"', level=1)
add_bullet("Select main branch, / (root) folder", level=1)
add_bullet("Click Save", level=1)

add_numbered("Your dashboard is now live at:")
add_code("https://YOURUSERNAME.github.io/REPONAME/dashboard.html")

# ══════════════════════════════════════════════════════════
# SECTION 7
# ══════════════════════════════════════════════════════════
doc.add_heading("7. How to Add or Remove Stocks", level=1)

doc.add_heading("To change the peer comparison stocks:", level=3)
doc.add_paragraph("Open fetch_data.py and find the PEERS section near the top:")
add_code('PEERS = {')
add_code('    "Urbanco":        "URBANCO.NS",')
add_code('    "Eternal":        "ETERNAL.NS",')
add_code('    "Swiggy":         "SWIGGY.NS",')
add_code('    ...etc...')
add_code('}')

doc.add_paragraph(
    "Add or remove lines as needed. The format is:"
)
add_code('    "Display Name":   "NSE_TICKER.NS",')

doc.add_paragraph(
    "To find the correct ticker symbol for any stock, search for it on "
    "nseindia.com and use the symbol shown in the URL."
)

# ══════════════════════════════════════════════════════════
# SECTION 8
# ══════════════════════════════════════════════════════════
doc.add_heading("8. Troubleshooting", level=1)

issues = [
    ("Dashboard shows old data",
     "Press Ctrl+Shift+R to force-refresh your browser. If still old, "
     "check GitHub Actions (see Section 5, Option A) to verify the daily "
     "job is running."),
    ("Day High / Day Low shows NaN",
     "This happens when the market hasn't closed yet or data isn't finalized. "
     "The dashboard automatically uses the last complete trading day's data. "
     "Check again after 8:30 PM."),
    ("Delivery % shows 0%",
     "NSE publishes delivery data after market close. If fetched too early, "
     "the dashboard estimates delivery % from historical patterns. The 8:30 PM "
     "schedule ensures data is usually available."),
    ("A stock in the peer table shows N/A",
     "The stock may be newly listed with less than 1 year of data, or the "
     "ticker symbol may have changed. Check nseindia.com for the current symbol."),
    ("GitHub Pages shows 404",
     "Go to repo Settings > Pages and make sure the source is set to "
     "'Deploy from a branch' with main / (root). Wait 2 minutes after saving."),
]

table3 = doc.add_table(rows=len(issues) + 1, cols=2)
table3.style = "Light Grid Accent 1"
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
table3.rows[0].cells[0].text = "Problem"
table3.rows[0].cells[1].text = "Solution"
for cell in table3.rows[0].cells:
    for p in cell.paragraphs:
        p.runs[0].bold = True
for i, (prob, sol) in enumerate(issues):
    table3.rows[i + 1].cells[0].text = prob
    table3.rows[i + 1].cells[1].text = sol

# ══════════════════════════════════════════════════════════
# SECTION 9
# ══════════════════════════════════════════════════════════
doc.add_heading("9. Quick Reference Card", level=1)

ref = [
    ("Dashboard URL", "https://ashishbansal-uc.github.io/urbanco-dashboard/dashboard.html"),
    ("Auto-update time", "8:30 PM IST, Monday-Friday"),
    ("GitHub repo", "https://github.com/ashishbansal-UC/urbanco-dashboard"),
    ("Manual refresh (web)", "GitHub repo > Actions > Run workflow"),
    ("Manual refresh (local)", "Double-click refresh.bat in the dashboard folder"),
    ("Force browser refresh", "Ctrl + Shift + R"),
    ("Data source", "Yahoo Finance (yfinance) + NSE API"),
    ("Stocks tracked", "Urbanco, Eternal, Swiggy, Meesho, Groww, Lenskart, Physics Wallah, Blackbuck, Pinelabs"),
]

table4 = doc.add_table(rows=len(ref) + 1, cols=2)
table4.style = "Light Grid Accent 1"
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
table4.rows[0].cells[0].text = "Item"
table4.rows[0].cells[1].text = "Details"
for cell in table4.rows[0].cells:
    for p in cell.paragraphs:
        p.runs[0].bold = True
for i, (item, detail) in enumerate(ref):
    table4.rows[i + 1].cells[0].text = item
    table4.rows[i + 1].cells[1].text = detail

doc.add_paragraph()
doc.add_paragraph()
add_para("Questions? Reach out to Ashish Bansal.", size=10,
         color=RGBColor(0x99, 0x99, 0x99), align=WD_ALIGN_PARAGRAPH.CENTER)

# ── Save ────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Guide saved to: {OUTPUT}")
